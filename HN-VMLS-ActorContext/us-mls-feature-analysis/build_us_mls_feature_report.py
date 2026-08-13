from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from pathlib import Path


OUTPUT = Path(__file__).resolve().parent / "Phan_tich_feature_MLS_My_2026.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
BORDER = "CBD3DD"
WHITE = "FFFFFF"
GREEN = "1F5D42"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def font_run(run, size=11, bold=False, italic=False, color="000000", name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    r_pr.append(sz)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    font_run(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(11)
        s.paragraph_format.left_indent = Inches(0.5)
        s.paragraph_format.first_line_indent = Inches(-0.25)
        s.paragraph_format.space_after = Pt(8)
        s.paragraph_format.line_spacing = 1.167

    if "Table Text" not in [s.name for s in styles]:
        s = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles["Table Text"]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(9)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(2)
    s.paragraph_format.line_spacing = 1.08

    if "Table Header" not in [s.name for s in styles]:
        s = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles["Table Header"]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(9)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.line_spacing = 1.05

    if "Source Note" not in [s.name for s in styles]:
        s = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        s = styles["Source Note"]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(9)
    s.font.color.rgb = RGBColor.from_string(MUTED)
    s.paragraph_format.space_before = Pt(4)
    s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.line_spacing = 1.0


def setup_section(section, first_page=False):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = first_page


def setup_running_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("FEATURE ANALYSIS  |  U.S. REAL ESTATE MLS")
    font_run(r, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Cập nhật 10/08/2026  |  ")
    font_run(r, size=9, color=MUTED)
    add_page_number(p)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        font_run(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        font_run(r)
    else:
        r = p.add_run(text)
        font_run(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    if level:
        p.paragraph_format.left_indent = Inches(0.75)
        p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    font_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    r = p.add_run(text)
    font_run(r)
    return p


def add_source_note(doc, text):
    p = doc.add_paragraph(style="Source Note")
    r = p.add_run(text)
    font_run(r, size=9, italic=True, color=MUTED)
    return p


def add_table(doc, headers, rows, widths, font_size=9, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, BLUE_GRAY)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Header"]
        if alignments and alignments[idx] == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        font_run(r, size=9, bold=True, color=NAVY)
    for row_idx, data in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(data):
            cell = cells[idx]
            set_cell_border(cell)
            if row_idx % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            if alignments and alignments[idx] == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            color = GREEN if str(text).startswith("5/") else (GOLD if str(text).startswith("3/") else "000000")
            font_run(r, size=font_size, bold=(idx == 0), color=color)
    return table


doc = Document()
configure_styles(doc)
setup_section(doc.sections[0], first_page=True)
setup_running_header_footer(doc.sections[0])

# Cover - editorial_cover pattern, with a restrained report-title override.
for _ in range(5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run("FEATURE ANALYSIS REPORT")
font_run(r, size=11, bold=True, color=GOLD)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Hệ thống MLS của Mỹ")
font_run(r, size=30, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Multiple Listing Service trong bất động sản")
font_run(r, size=16, color=DARK_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(54)
r = p.add_run("Cấu trúc thị trường, capability map, dữ liệu, compliance và roadmap sản phẩm")
font_run(r, size=11, italic=True, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Phạm vi: Residential MLS cấp địa phương/khu vực tại Hoa Kỳ")
font_run(r, size=10.5, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run("Phiên bản: 1.0  |  Cập nhật: 10/08/2026")
font_run(r, size=10, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(58)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Lưu ý: MLS trong báo cáo là Multiple Listing Service, không phải Major League Soccer.")
font_run(r, size=9.5, italic=True, color=MUTED)

doc.add_page_break()

doc.add_heading("Tóm tắt điều hành", level=1)
add_body(doc, "MLS của Mỹ nên được hiểu là một hạ tầng thị trường có quy tắc, không chỉ là website tìm nhà. Nó đồng thời thực hiện ba việc: tạo kho dữ liệu listing dùng chung giữa các brokerage; điều phối hợp tác và quyền truy cập giữa các thành viên; phân phối dữ liệu có kiểm soát sang website, ứng dụng và hệ thống nghiệp vụ. NAR mô tả khoảng 500 MLS tại Mỹ; RESO ghi nhận 489 hệ thống đang hoạt động vào tháng 7/2026. [1][2]")
add_body(doc, "Kết luận chính: năng lực lõi của MLS đã rất trưởng thành ở quản lý listing, tìm kiếm bản đồ, saved search/alert, CMA, dữ liệu lịch sử, phân phối IDX và kiểm soát tuân thủ. Điểm yếu mang tính cấu trúc nằm ở phân mảnh khu vực, khác biệt rule/data giữa các MLS, trải nghiệm người dùng không đồng nhất và việc nhiều quy trình giao dịch phải dựa vào sản phẩm tích hợp bên ngoài.")
add_body(doc, "Thay đổi chính sách từ 2024-2026 tạo ra yêu cầu sản phẩm mới: không được hiển thị offer of compensation trên MLS; buyer agreement phải hoàn tất trước khi touring trong phạm vi chính sách áp dụng; và listing có thể ở trạng thái delayed marketing, vẫn thấy trong MLS nhưng chưa được đưa ra IDX/syndication. Vì vậy workflow, quyền phân phối, audit trail và rule engine phải là capability lõi, không phải phần cấu hình phụ. [8][9][10]")

add_table(doc,
          ["Kết luận", "Ý nghĩa sản phẩm"],
          [
              ("MLS là market infrastructure", "Ưu tiên data integrity, quyền truy cập, audit và rule enforcement ngang với UX."),
              ("Không có một MLS quốc gia duy nhất", "Thiết kế multi-tenant, local rules, schema mapping và cross-market identity ngay từ đầu."),
              ("Front-end chỉ là một lớp", "Core data, licensing/API, compliance và distribution có giá trị chiến lược hơn một màn hình search đẹp."),
              ("AI mới ở giai đoạn bổ sung", "Dùng AI để tăng tốc nhập liệu, tìm kiếm và phát hiện lỗi; quyết định compliance vẫn phải giải thích được."),
          ],
          [2400, 6960], font_size=9.3)
add_source_note(doc, "Nguồn tổng hợp: NAR, RESO, CRMLS, ICE và MLS Grid; xem Phụ lục nguồn.")

doc.add_heading("1. Phạm vi và cách đọc báo cáo", level=1)
add_body(doc, "Báo cáo phân tích mẫu hình chung của residential MLS tại Mỹ, sử dụng tài liệu chính sách NAR, tiêu chuẩn RESO và chức năng thực tế từ các hệ thống/nhà vận hành lớn. Đây không phải review của một sản phẩm duy nhất. Một MLS địa phương có thể dùng Matrix, Paragon, Flexmls, Perchwell hoặc hệ thống khác; bundle tích hợp cũng khác nhau.")
add_bullet(doc, "Đối tượng đọc: product owner, founder, business analyst, solution architect và đội vận hành thị trường bất động sản.")
add_bullet(doc, "Tiêu chí đánh giá: giá trị người dùng, độ trưởng thành thị trường, độ khó dữ liệu/compliance và mức độ nên đưa vào MVP.")
add_bullet(doc, "Ngoài phạm vi: luật chi tiết từng bang, commercial MLS/CIE chuyên sâu, title/escrow/mortgage servicing và tư vấn pháp lý.")

doc.add_heading("2. MLS là gì và hệ sinh thái vận hành ra sao?", level=1)
add_body(doc, "Về bản chất, MLS là nền tảng online tổng hợp listing từ nhiều brokerage trong một thị trường và cung cấp dữ liệu để môi giới tìm, chia sẻ, phân tích và quảng bá tài sản. Giá trị mạng lưới đến từ việc các đối thủ cạnh tranh cùng đóng góp dữ liệu theo rule chung, tạo inventory đủ rộng và lịch sử đủ tin cậy cho buyer, seller, appraiser và lender. [1]")
add_table(doc,
          ["Lớp", "Chủ thể tiêu biểu", "Vai trò"],
          [
              ("Governance", "MLS operator, REALTOR® association, broker-owned MLS", "Ban hành local rules, membership, enforcement, data licenses và vendor policy."),
              ("Core platform", "Matrix, Paragon, Flexmls, Perchwell...", "Listing database, add/edit, search, portals, CMA, reporting và admin console."),
              ("Data standard", "RESO Data Dictionary + Web API", "Chuẩn hóa resource/field/lookup và cách vận chuyển dữ liệu giữa MLS, broker, vendor."),
              ("Workflow ecosystem", "Showing, lockbox, forms, transaction, tax, CRM", "Bổ sung tác vụ trước và sau khi listing được tìm thấy."),
              ("Distribution", "IDX, VOW, syndication, back-office feeds", "Đưa dữ liệu tới website brokerage, portal, app và hệ thống nội bộ theo quyền sử dụng."),
              ("Consumer surface", "Broker/agent site, app, client portal, national portal", "Trải nghiệm tìm kiếm và cộng tác cho buyer/seller; thường là downstream của MLS."),
          ],
          [1450, 3000, 4910], font_size=8.8)
add_source_note(doc, "NAR ước tính khoảng 500 MLS; RESO ghi nhận 489 hệ thống đang hoạt động và ít nhất 90% có Web API được RESO chứng nhận tại thời điểm cập nhật. [1][2]")

doc.add_heading("3. Nhóm người dùng và Jobs-to-be-Done", level=1)
add_table(doc,
          ["Persona", "Jobs-to-be-Done chính", "Kết quả mong đợi"],
          [
              ("Listing agent", "Tạo listing đúng rule; quản lý media, status, open house, showing và exposure.", "Đăng nhanh, ít lỗi, phủ thị trường rộng và theo dõi hiệu quả."),
              ("Buyer agent", "Tìm inventory phù hợp; theo dõi thay đổi; gửi shortlist; đặt lịch xem; làm CMA.", "Không bỏ lỡ listing, tư vấn nhanh và có bằng chứng dữ liệu."),
              ("Broker/office admin", "Giám sát roster, listing, quyền, feed, compliance và hiệu suất đội ngũ.", "Giảm rủi ro và có visibility toàn văn phòng."),
              ("MLS compliance/admin", "Kiểm tra dữ liệu, xử lý vi phạm, cấu hình rule, phê duyệt vendor/data license.", "Dữ liệu đầy đủ, chính xác, có audit trail và enforcement nhất quán."),
              ("Buyer/seller", "Nhận listing phù hợp, phản hồi, xem lịch sử/so sánh và hiểu chiến lược thị trường.", "Thông tin kịp thời, rõ nguồn và làm việc hiệu quả với agent."),
              ("Appraiser/lender/vendor", "Truy cập dữ liệu lịch sử, comps hoặc feed được cấp phép.", "Dữ liệu chuẩn hóa, cập nhật và đúng phạm vi sử dụng."),
          ],
          [1750, 4200, 3410], font_size=8.8)

doc.add_heading("4. Luồng nghiệp vụ end-to-end", level=1)
for step in [
    "Onboarding & entitlement: broker tham gia MLS; agent/subscriber được gắn office, license, role, product entitlement và SSO/MFA.",
    "Create listing: chọn property type, hợp đồng/listing service, nhập địa chỉ/APN, auto-fill public record, geocode, media và remarks.",
    "Validate & publish: rule engine kiểm tra field bắt buộc, logic chéo, media, thời hạn nộp; xác định Active, Coming Soon, Office Exclusive hoặc Delayed Marketing.",
    "Discover & collaborate: buyer agent search theo field/map, lưu search, nhận alert/hotsheet, gửi portal/email và theo dõi phản hồi.",
    "Show & qualify: tích hợp showing availability/request, lockbox access, open house, financing calculator và tài liệu liên quan.",
    "Offer/contract: cập nhật trạng thái pending/contingent theo local rules; hand-off sang forms, e-sign và transaction management.",
    "Close & retain history: cập nhật closed price/date, concessions nếu local MLS cho phép, buyer office/member và lịch sử thay đổi.",
    "Distribute & analyze: đồng bộ IDX/VOW/syndication/back-office; tạo CMA, market statistics, agent/broker performance và compliance reports.",
]:
    add_number(doc, step)

doc.add_heading("5. Capability map và feature analysis", level=1)
add_body(doc, "Thang điểm trưởng thành 1-5 phản ánh mức phổ biến và mức hoàn thiện quan sát được trên các MLS lớn, không phải điểm của mọi local MLS. CRMLS hiện so sánh đồng thời Paragon, Matrix, Flex và Perchwell với tập feature gồm map search, add/edit listing, auto-email, CMA, contact, tax, help/compliance và mobile. [5]")
add_table(doc,
          ["Capability", "Feature cốt lõi", "Mức"],
          [
              ("Listing management", "Draft/copy, smart fields, media, geocode, public-record autofill, status/change, open house.", "5/5"),
              ("Search & map", "Quick/advanced/cross-property, polygon/radius, exclude area, map result, MLS#/agent/office/open house.", "5/5"),
              ("Saved search & alerts", "Auto-email, concierge approval, hotsheet, change alerts, expiry, client view tracking.", "5/5"),
              ("Client collaboration", "Portal, favorite/reject/comment, share, activity, buyer/seller reports.", "4/5"),
              ("CMA & valuation", "Comparable search, adjustments, branded CMA, seller report, market statistics.", "4/5"),
              ("Showing & access", "Availability, request/confirm, feedback, open house, lockbox assignment/access.", "4/5"),
              ("Distribution", "IDX, VOW, syndication, broker back-office, opt-in/out, refresh, attribution.", "5/5"),
              ("Data/API", "RESO schema, Web API/OData, replication/delta, media download, licenses and payload rules.", "4/5"),
              ("Compliance", "Validation, deadlines, violation report, notices, audit, fines/work queues, rule configuration.", "5/5"),
              ("Roster & identity", "Member/office/team/license, SSO, roles, entitlements, reciprocal access.", "4/5"),
              ("Transaction workflow", "Forms, e-sign, disclosures, earnest money, transaction room via integrations.", "3/5"),
              ("CRM & productivity", "Contacts, groups, tasks, reverse prospecting, activity and basic lead management.", "3/5"),
              ("Mobile", "Responsive search, near-me, listing maintenance, contacts, CMA and alerts.", "4/5"),
              ("AI assistance", "Natural-language search, description draft, photo enhancement, anomaly detection.", "2/5"),
              ("Cross-market interoperability", "Data shares, reciprocal links, common feed/license; vẫn bị chia cắt theo membership/rule.", "3/5"),
          ],
          [2100, 6500, 760], font_size=8.4, alignments=["left", "left", "center"])

doc.add_heading("5.1 Listing management", level=2)
add_body(doc, "Đây là system of record của listing. Bộ feature trưởng thành gồm lưu bản nháp, copy listing, tự điền từ public/tax record, geocode, smart field theo property type, kiểm tra lỗi, quản lý nhiều ảnh/tài liệu/virtual tour, sắp xếp ảnh, caption và cập nhật nhanh status/price. CRMLS cho thấy các capability này xuất hiện rộng trên bốn nền tảng chính; Paragon Connect bổ sung speech-to-text và error tracking trong EasyListing. [5][7]")
add_bullet(doc, "Điểm mạnh: dữ liệu có cấu trúc sâu, lịch sử thay đổi và kiểm soát theo role/office.")
add_bullet(doc, "Pain point: form dài, local field khác nhau, nhiều validation chỉ xuất hiện cuối quy trình và duplicate property/listing.")
add_bullet(doc, "Opportunity: progressive form, pre-fill có provenance, validation theo thời gian thực, duplicate detection và checklist theo status.")

doc.add_heading("5.2 Search, map và inventory monitoring", level=2)
add_body(doc, "Search là feature hằng ngày có độ trưởng thành cao: exact/contains, field filter, address/APN/MLS number, agent/office, property type, open house, quick/cross-property/CMA search; map hỗ trợ polygon, radius và vùng loại trừ. Hotsheet và Market Watch tập trung vào listing mới, price/status change và listing sắp hết hạn. [5]")
add_bullet(doc, "Điểm mạnh: độ chi tiết field cao hơn đa số consumer portal và hỗ trợ saved display/report.")
add_bullet(doc, "Pain point: cognitive load lớn; tên field và lookup phụ thuộc local MLS; query khó chuyển giữa thị trường.")
add_bullet(doc, "Opportunity: natural-language-to-filter có phần giải thích, query templates theo use case và unified search ontology.")

doc.add_heading("5.3 Prospecting và client collaboration", level=2)
add_body(doc, "Agent lưu search gắn với contact, cấu hình lịch auto-email hoặc concierge review, sau đó theo dõi view/favorite/comment trong portal. Một số hệ thống có reverse prospecting: listing agent nhìn thấy buyer criteria khớp mà không lộ danh tính consumer. RESO 2.0 đã mô hình hóa Contacts, SavedSearch, Prospecting, ContactListings và ContactListingNotes, cho thấy collaboration đang trở thành dữ liệu chuẩn chứ không chỉ UI. [3][5]")
add_bullet(doc, "Khoảng trống: portal thường thiên về email feed, chưa phải workspace giao dịch đa bên; mobile consumer UX không đồng nhất.")
add_bullet(doc, "Opportunity: shared shortlist, explainable match score, tour planner, notification preference và consent management.")

doc.add_heading("5.4 CMA, statistics và advisory", level=2)
add_body(doc, "CMA kết hợp tìm comparable, điều chỉnh, lịch sử listing/sale, tax/public record và template báo cáo để agent tư vấn giá. Market statistics bổ sung inventory, median price, days on market, absorption và trend theo geography/property type. Nhiều MLS bundle sản phẩm như Cloud CMA, InfoSparks hoặc RPR thay vì tự xây mọi lớp analytics. [6]")
add_bullet(doc, "Điểm mạnh: dữ liệu closed/listing history tạo lợi thế so với search portal công khai.")
add_bullet(doc, "Rủi ro: comp selection bias, khác biệt non-disclosure state, measurement/source không đồng nhất.")
add_bullet(doc, "Opportunity: provenance cho từng dữ liệu, adjustment rationale, confidence band và reproducible CMA snapshot.")

doc.add_heading("5.5 Showing, lockbox và open house", level=2)
add_body(doc, "Showing thường là integration nhưng nằm trong journey lõi: seller availability, request/confirmation, instruction bảo mật, feedback và audit. Lockbox (ví dụ Supra/SentriLock) được gắn vào listing; open house có public/broker/office type. RESO đã chuẩn hóa Showing, ShowingRequest, ShowingAppointment, ShowingAvailability và LockOrBox. [3][5]")
add_bullet(doc, "Yêu cầu quan trọng: tách public instructions khỏi confidential access instructions; kiểm soát role, thời gian và log mở khóa.")

doc.add_heading("5.6 Distribution: IDX, VOW, syndication và back office", level=2)
add_body(doc, "IDX cho phép participant hiển thị aggregated listing trên website/app được kiểm soát; VOW dành cho dịch vụ brokerage online sau khi thiết lập quan hệ broker-consumer; syndication đưa listing tới portal phân phối; back-office feed phục vụ CRM, accounting, analytics hoặc broker technology. NAR yêu cầu IDX download phản ánh update/status change ít nhất mỗi 12 giờ, đồng thời cho phép MLS yêu cầu firewall và audit trail. [11][12]")
add_bullet(doc, "Quyền phân phối cần độc lập với trạng thái listing: một listing có thể thấy trong MLS nhưng chưa được ra IDX/syndication.")
add_bullet(doc, "Opt-out, attribution, sold-data display, field suppression, refresh/deletion và downstream vendor access phải là policy objects có version.")
add_bullet(doc, "MLS Grid minh họa mô hình gom nhiều MLS về một feed, một license và một compliance process, nhưng broker chỉ nhận dữ liệu từ MLS nơi họ có participatory rights. [13]")

doc.add_heading("5.7 Compliance, admin và data quality", level=2)
add_body(doc, "Một MLS có giá trị vì dữ liệu được đóng góp theo deadline, field rule, status rule và quyền sử dụng chung. Admin cần rule configuration, validation severity, violation report, anonymous report, work queue, notice, cure period, fine/waiver, audit log, data-license approval và export chứng cứ. Local rule variability khiến compliance-as-code quan trọng hơn hard-code trong form.")
add_bullet(doc, "Rule nên có scope theo property type, listing status, organization, effective date và distribution channel.")
add_bullet(doc, "Mọi override cần lý do, người phê duyệt, timestamp và before/after values.")
add_bullet(doc, "Dashboard chất lượng nên theo dõi completeness, timeliness, duplicate, stale status, invalid geo/media và downstream sync failures.")

doc.add_heading("6. Mô hình dữ liệu và kiến trúc tích hợp", level=1)
add_body(doc, "RESO Data Dictionary là ontology dùng chung cho resource, field và lookup; Web API là lớp vận chuyển hiện đại dựa trên công nghệ web/OData. Data Dictionary 2.0 công bố 41 resource và 1.745 field, bao gồm Property, Member, Office, Media, OpenHouse, Contacts, SavedSearch, Prospecting, Showing, LockOrBox, transactional history và internet tracking. [3][4]")
add_table(doc,
          ["Domain", "Entity/resource trọng tâm", "Yêu cầu thiết kế"],
          [
              ("Inventory", "Property, Media, OpenHouse, HistoryTransactional", "Stable IDs, version history, status lifecycle, media ordering và provenance."),
              ("Identity", "Member, Office, Team, Association, licenses, OUID", "Multi-MLS identity, role/entitlement, license effective dates và merge rules."),
              ("Demand", "Contact, SavedSearch, Prospecting, ContactListing", "Consent, preference, match events, activity retention và broker ownership."),
              ("Access", "Showing, ShowingRequest/Appointment/Availability, LockOrBox", "Time-bound authorization, confidential fields và immutable audit."),
              ("Distribution", "Payload, Rules, Queue/EntityEvent, InternetTracking", "Channel-level policy, delta/event, delete/tombstone, retry và usage analytics."),
          ],
          [1550, 3350, 4460], font_size=8.8)

doc.add_heading("6.1 Kiến trúc tham chiếu", level=2)
for item in [
    "Policy & identity plane: organization, roster, license, roles, entitlements, rule versions, data-use agreements.",
    "Canonical MLS data plane: normalized resources, stable keys, effective-dated listing lifecycle, media store, geo index và history.",
    "Workflow plane: add/edit, search, portal, CMA, showing, compliance workbench và admin console.",
    "Integration plane: RESO Web API, delta/event queue, webhooks, batch replication, SSO, tax/public records, lockbox, forms và transaction vendors.",
    "Distribution plane: payload builder theo IDX/VOW/back-office/syndication, field suppression, attribution, watermark/cache policy và monitoring.",
    "Analytics plane: operational metrics, data quality, market statistics, usage, vendor compliance và model/AI audit.",
]:
    add_bullet(doc, item)

doc.add_heading("6.2 API và đồng bộ", level=2)
add_body(doc, "RESO Web API giảm mapping riêng lẻ, nhưng chứng nhận ngoài không đồng nghĩa mọi core database đã chuẩn hóa hoàn toàn. Một số MLS map schema nội bộ sang RESO ở lớp API. MLS Grid dùng OAuth 2 token và replication-oriented API; consumer phải lưu hoặc cache media đúng điều khoản thay vì hotlink tùy ý. [2][4][13]")
add_bullet(doc, "Khuyến nghị: hỗ trợ cursor/delta theo modification timestamp hoặc event sequence, tombstone cho delete/withdrawn, idempotency và reconciliation job.")
add_bullet(doc, "SLO tối thiểu cần tách: thời gian cập nhật core, thời gian publish từng channel, media availability và freshness downstream.")

doc.add_heading("7. Chính sách và compliance ảnh hưởng trực tiếp tới feature", level=1)
add_table(doc,
          ["Chính sách", "Yêu cầu sản phẩm", "Tác động UX/data"],
          [
              ("Clear Cooperation", "Listing được public market phải nộp vào MLS trong 1 business day, tùy phạm vi rule áp dụng.", "Deadline clock, proof of submission, public-marketing trigger và exception workflow. [8]"),
              ("Office Exclusive", "Nộp cho MLS nhưng không disseminate tới participant; không public market.", "Seller certification, restricted visibility, access log và no-publication guardrail. [9]"),
              ("Delayed Marketing", "Listing thấy trong MLS nhưng tạm hoãn IDX/syndication theo thời hạn local MLS.", "Channel-specific availability, timer, seller disclosure và auto-release. [9]"),
              ("No compensation in MLS", "Không có offer of compensation trong field, remarks, data feed hoặc filter.", "Content scanning, prohibited-field policy, API suppression; đàm phán diễn ra off-MLS. [10]"),
              ("Written buyer agreement", "Participant working with buyer phải ký agreement trước touring, trừ khi luật áp dụng quy định khác.", "Pre-tour attestation/workflow, document link và broker supervision; không nhất thiết MLS giữ bản copy. [10]"),
              ("IDX policy", "Consent/opt-out, display rules, security, audit và refresh tối thiểu 12 giờ.", "License-aware payload, attribution, monitoring và revocation. [11]"),
          ],
          [1700, 3800, 3860], font_size=8.4)
add_source_note(doc, "NAR policies chủ yếu áp dụng cho REALTOR® association MLS và các participant trong phạm vi chính sách/settlement; local MLS và luật bang có thể thêm yêu cầu. Báo cáo không thay thế tư vấn pháp lý.")

doc.add_heading("8. Đánh giá strengths, weaknesses và khoảng trống", level=1)
doc.add_heading("8.1 Điểm mạnh", level=2)
add_bullet(doc, "Network effect: inventory do nhiều brokerage cùng đóng góp, giúp một agent phục vụ buyer trên tập listing rộng.")
add_bullet(doc, "Data depth: field chi tiết, agent-only remarks, history và sold/comparable data có giá trị nghiệp vụ cao.")
add_bullet(doc, "Accountability: listing gắn member/office/license; change và violation có thể audit.")
add_bullet(doc, "Distribution control: một nguồn listing có thể cấp nhiều payload với rule khác nhau.")
add_bullet(doc, "Standards momentum: RESO Web API và Data Dictionary giảm chi phí tích hợp, dù chưa xóa hết local variance.")

doc.add_heading("8.2 Điểm yếu cấu trúc", level=2)
add_bullet(doc, "Fragmentation: agent/broker đa thị trường có thể phải trả phí, đăng nhập và tuân thủ nhiều bộ rule khác nhau.")
add_bullet(doc, "Inconsistent schemas/statuses: RESO giúp chuẩn hóa exchange nhưng local fields, lookup và business semantics vẫn khác.")
add_bullet(doc, "Legacy UX: form dày, desktop-first, nhiều module rời và training overhead lớn.")
add_bullet(doc, "Integration dependency: showing, transaction, CRM, tax, analytics và media thường đến từ vendor bundle; trải nghiệm đứt đoạn.")
add_bullet(doc, "Latency/compliance burden: feed replication, cache và vendor licensing làm tăng độ trễ, chi phí giám sát và rủi ro misuse.")
add_bullet(doc, "AI governance gap: feature AI xuất hiện không đồng đều; cần kiểm soát hallucination, fair housing, source attribution và human approval.")

doc.add_heading("9. Cơ hội sản phẩm ưu tiên", level=1)
add_table(doc,
          ["Ưu tiên", "Cơ hội", "Giá trị / acceptance signal"],
          [
              ("P0", "Policy-aware listing lifecycle", "Status và channel rights độc lập; rule effective-dated; 100% transition có audit."),
              ("P0", "Quality-by-design add/edit", "Pre-fill có nguồn, inline validation, duplicate check; giảm correction/violation và time-to-publish."),
              ("P0", "Entitlement & data licensing", "Một identity graph cho member/office/team/vendor; revoke có hiệu lực xuyên API/channel."),
              ("P0", "Reliable distribution platform", "IDX/VOW/back-office payload cấu hình; delta/tombstone/retry; đo end-to-end freshness."),
              ("P1", "Unified search + explainable AI", "Natural language tạo filter nhìn thấy được; agent có thể chỉnh và lưu; zero hidden steering."),
              ("P1", "Consumer collaboration workspace", "Shortlist, comment, tour plan, consent/notification; tăng engagement có thể đo."),
              ("P1", "Compliance workbench", "Rule severity, case queue, evidence, cure/fine/appeal và KPI chất lượng."),
              ("P1", "Cross-market normalization", "Common IDs, dedupe, schema mapping và query portability giữa MLS được cấp quyền."),
              ("P2", "Explainable CMA/analytics", "Snapshot tái tạo được, provenance, adjustment rationale, confidence và bias checks."),
              ("P2", "Ecosystem marketplace", "Sandbox, certification, scoped OAuth, usage billing/analytics và vendor scorecard."),
          ],
          [900, 3020, 5440], font_size=8.6)

doc.add_heading("10. Đề xuất phạm vi nếu xây một MLS tương đương", level=1)
doc.add_heading("10.1 MVP - 0 đến 6 tháng", level=2)
add_bullet(doc, "Core identity: organization, office, member, role, license, SSO/MFA, entitlement.")
add_bullet(doc, "Canonical Property/Listing/Media/OpenHouse model; lifecycle và history; configurable required fields/lookups.")
add_bullet(doc, "Add/edit có draft, copy, public-record/geocode integration, photo ordering và validation.")
add_bullet(doc, "Advanced field search + map polygon/radius; saved search; hotsheet; result/report export.")
add_bullet(doc, "IDX-style feed đầu tiên qua RESO-aligned API; channel consent, attribution, refresh, revoke và audit.")
add_bullet(doc, "Admin/compliance cơ bản: rule versioning, violation report, work queue và data-quality dashboard.")

doc.add_heading("10.2 V1 - 6 đến 12 tháng", level=2)
add_bullet(doc, "Contacts, auto-email/concierge, client portal, favorite/comment và activity.")
add_bullet(doc, "CMA, comparable workflow, basic market statistics và seller/listing activity report.")
add_bullet(doc, "Showing request/availability, open house và lockbox integrations.")
add_bullet(doc, "VOW/back-office/syndication payloads, vendor onboarding, usage monitoring và multi-MLS mapping.")
add_bullet(doc, "Mobile responsive parity cho search, alerts, listing edit và CMA-lite.")

doc.add_heading("10.3 V2 - 12 đến 24 tháng", level=2)
add_bullet(doc, "Cross-market identity, shared listing dedupe, reciprocal/data-share workflows và unified license experience.")
add_bullet(doc, "Event-driven distribution, near-real-time consumer alerts và observability xuyên downstream.")
add_bullet(doc, "Explainable AI search, description draft, photo/data anomaly detection và compliance triage với human approval.")
add_bullet(doc, "Broker analytics, vendor marketplace, sandbox/certification và configurable product bundles.")

doc.add_heading("11. KPI và non-functional requirements", level=1)
add_table(doc,
          ["Dimension", "KPI gợi ý", "Mục tiêu thiết kế"],
          [
              ("Freshness", "P50/P95 core-to-channel latency; stale listing count", "Đo riêng từng payload; alert khi vượt SLO; hỗ trợ replay/reconciliation."),
              ("Data quality", "Completeness, correction rate, duplicate rate, geocode/media error", "Validation trước publish, provenance và issue ownership rõ."),
              ("Search", "P95 query latency, zero-result rate, saved-search match delay", "Geo index + structured filter; explainable query plan cho AI search."),
              ("Reliability", "API availability, event lag, failed sync, recovery time", "Idempotent processing, backpressure, retry, tombstone và disaster recovery."),
              ("Compliance", "Violation/1.000 listing, time-to-cure, repeat rate", "Rule effective date, evidence bundle, role separation và immutable audit."),
              ("Security", "MFA coverage, suspicious access, revoke latency", "Least privilege, scoped token, confidential field masking và device/session controls."),
              ("Adoption", "Weekly active agents, time-to-first-search/listing, portal engagement", "Role-based onboarding, embedded help và workflow analytics."),
          ],
          [1400, 3520, 4440], font_size=8.8)

doc.add_heading("12. Kết luận", level=1)
add_body(doc, "MLS Mỹ thành công không phải vì một feature đơn lẻ, mà vì tổ hợp của network participation, dữ liệu có cấu trúc, rule enforcement và khả năng phân phối. Nếu chỉ sao chép search portal, sản phẩm sẽ bỏ lỡ phần khó và có giá trị nhất: identity/entitlement, lifecycle, data quality, licensing, compliance và interoperability.")
add_body(doc, "Chiến lược hợp lý là xây một core MLS có policy engine và data platform vững, sau đó mở rộng thành ecosystem. UX hiện đại và AI có thể tạo lợi thế, nhưng phải đứng trên dữ liệu có provenance, filter giải thích được, quyền truy cập minh bạch và audit end-to-end.")

doc.add_page_break()
doc.add_heading("Phụ lục A - Glossary", level=1)
add_table(doc,
          ["Thuật ngữ", "Định nghĩa ngắn"],
          [
              ("MLS", "Multiple Listing Service - hạ tầng chia sẻ listing và dữ liệu thị trường giữa các brokerage theo rule chung."),
              ("IDX", "Internet Data Exchange - data use cho phép participant hiển thị aggregated MLS listing trên website/app được kiểm soát."),
              ("VOW", "Virtual Office Website - trải nghiệm brokerage online sau khi thiết lập quan hệ broker-consumer theo quy định."),
              ("Syndication", "Phân phối listing tới portal/app bên thứ ba theo quyết định của broker/seller và local rules."),
              ("CMA", "Comparative Market Analysis - phân tích tài sản so sánh để tư vấn giá/listing strategy."),
              ("RESO", "Real Estate Standards Organization - tổ chức xây dựng tiêu chuẩn dữ liệu và API cho bất động sản."),
              ("Payload", "Tập field/record được phép cấp cho một use case cụ thể như IDX, VOW hoặc back office."),
              ("Hotsheet", "Danh sách thay đổi mới gần đây: new listing, price/status change, back on market, expiring..."),
              ("Reverse prospecting", "Listing agent thấy mức độ khớp với saved-search demand mà không cần lộ danh tính consumer."),
          ],
          [1900, 7460], font_size=9)

doc.add_heading("Phụ lục B - Nguồn tham khảo", level=1)
sources = [
    ("[1] NAR - Consumer Guide: Multiple Listing Services (MLSs)", "https://www.nar.realtor/node/200359"),
    ("[2] RESO - Certification and MLS Map", "https://www.reso.org/certification/"),
    ("[3] RESO - Data Dictionary 2.0", "https://dd.reso.org/DD2.0/"),
    ("[4] RESO - RESO Web API", "https://www.reso.org/reso-web-api/"),
    ("[5] CRMLS - 2026 MLS Systems Comparison Chart", "https://kb.crmls.org/wp-content/uploads/2025/01/2026_MLS_Systems_Comparison_Chart.pdf"),
    ("[6] CRMLS - 2026 Product Solutions Matrix", "https://go.crmls.org/wp-content/uploads/2026/05/2026_Product_Solutions_Matrix.pdf"),
    ("[7] ICE Mortgage Technology - Paragon Connect MLS Platform", "https://mortgagetech.ice.com/products/paragon-connect-mls-platform"),
    ("[8] NAR - MLS Clear Cooperation Policy", "https://www.nar.realtor/about-nar/policies/mls-clear-cooperation-policy"),
    ("[9] NAR - Multiple Listing Options for Sellers", "https://www.nar.realtor/about-nar/policies/multiple-listing-options-for-sellers"),
    ("[10] NAR - No Compensation Offers in MLS; Written Buyer Agreements", "https://www.nar.realtor/handbook-on-multiple-listing-policy/no-compensation-offers-in-mls-section-4-written-buyer-agreements-required-policy-statement-8-13"),
    ("[11] NAR - Internet Data Exchange (IDX) Policy", "https://www.nar.realtor/handbook-on-multiple-listing-policy/advertising-print-and-electronic-section-1-internet-data-exchange-idx-policy-policy-statement-7-58"),
    ("[12] NAR - Virtual Office Websites Policy", "https://www.nar.realtor/handbook-on-multiple-listing-policy/virtual-office-websites-policy-governing-use-of-mls-data-in-connection-with-internet-brokerage"),
    ("[13] MLS Grid - Documentation Overview", "https://docs.mlsgrid.com/"),
    ("[14] CRMLS - Introduction to IDX and Listing Distribution Options", "https://go.crmls.org/wp-content/uploads/2026/03/2026_An_Introduction_to_IDX_and_Your_Listing_Distribution_Options.pdf"),
]
for label, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.05
    add_hyperlink(p, label, url)

add_source_note(doc, "Truy cập và đối chiếu ngày 10/08/2026. Các policy/local rule có thể tiếp tục thay đổi; cần kiểm tra MLS và luật bang cụ thể trước khi triển khai.")

# Core properties and document metadata.
props = doc.core_properties
props.title = "Feature Analysis Report - Hệ thống MLS của Mỹ"
props.subject = "Multiple Listing Service trong bất động sản Hoa Kỳ"
props.author = "OpenAI Codex"
props.keywords = "MLS, Multiple Listing Service, feature analysis, RESO, IDX, VOW, real estate"
props.comments = "Prepared from official NAR, RESO, CRMLS, ICE and MLS Grid sources; updated 2026-08-10."

# Keep tables from splitting header rows and prevent accidental fixed row heights.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

doc.save(OUTPUT)
print(OUTPUT)
